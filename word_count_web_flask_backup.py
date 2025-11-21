import os
import csv
import docx
from flask import Flask, render_template, request, jsonify, send_file, Response
import io
from PyPDF2 import PdfReader
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

app = Flask(__name__)
# 设置最大请求体大小为 100MB，防止单次请求过大
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

def get_word_count(file_path):
    """
    读取 docx 文件并统计字数 (统计逻辑：纯文本字符数，去除空格和换行)
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text_content = "".join(full_text)
        clean_text = text_content.replace(" ", "").replace("\n", "").replace("\t", "").replace("\r", "")
        return len(clean_text), "成功"
    except Exception as e:
        return 0, f"失败: {str(e)}"

def get_pdf_word_count(file_path):
    """
    读取 PDF 文件并统计字数 (统计逻辑：纯文本字符数，去除空格和换行)
    """
    try:
        reader = PdfReader(file_path)
        full_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
        text_content = "".join(full_text)
        clean_text = text_content.replace(" ", "").replace("\n", "").replace("\t", "").replace("\r", "")
        return len(clean_text), "成功"
    except Exception as e:
        return 0, f"失败: {str(e)}"

def get_txt_word_count(file_path):
    """
    读取 TXT 文件并统计字数 (统计逻辑：纯文本字符数，去除空格和换行)
    支持多种编码：UTF-8, GBK, GB2312
    """
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text_content = f.read()
            clean_text = text_content.replace(" ", "").replace("\n", "").replace("\t", "").replace("\r", "")
            return len(clean_text), "成功"
        except (UnicodeDecodeError, LookupError):
            continue

    # 如果所有编码都失败
    return 0, "失败: 无法识别文件编码"

def get_word_count_unified(file_path):
    """
    统一的字数统计入口，根据文件扩展名分发到对应的处理函数
    支持 .docx, .pdf, .txt, .md 文件
    """
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == '.docx':
        return get_word_count(file_path)
    elif file_extension == '.pdf':
        return get_pdf_word_count(file_path)
    elif file_extension in ['.txt', '.md']:
        return get_txt_word_count(file_path)
    else:
        return 0, f"失败: 不支持的文件格式 {file_extension}"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/analyze', methods=['POST'])
def analyze():
    data = request.json
    folder_path = data.get('folder_path', '').strip()

    # Remove quotes if user dragged and dropped folder
    if (folder_path.startswith('"') and folder_path.endswith('"')) or (folder_path.startswith("'") and folder_path.endswith("'")):
        folder_path = folder_path[1:-1]

    if not folder_path or not os.path.isdir(folder_path):
        return jsonify({'error': '无效的文件夹路径，请检查后重试'}), 400

    # 支持 .docx, .pdf, .txt, .md 文件
    supported_extensions = ('.docx', '.pdf', '.txt', '.md')
    supported_files = [f for f in os.listdir(folder_path)
                       if f.lower().endswith(supported_extensions) and not f.startswith('~$')]

    if not supported_files:
        return jsonify({'error': '该文件夹下没有找到支持的文件 (.docx, .pdf, .txt, .md)'}), 404

    results = []
    for filename in supported_files:
        file_path = os.path.join(folder_path, filename)
        char_count, status = get_word_count_unified(file_path)
        file_type = os.path.splitext(filename)[1].lower()  # 获取文件扩展名
        results.append({
            'filename': filename,
            'file_type': file_type,
            'char_count': char_count,
            'status': status
        })

    # Sort by filename
    results.sort(key=lambda x: x['filename'])

    return jsonify({'results': results, 'count': len(results)})

@app.route('/api/analyze_upload', methods=['POST'])
def analyze_upload():
    if 'files[]' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    files = request.files.getlist('files[]')
    if not files:
        return jsonify({'error': '未选择文件'}), 400

    results = []

    # Create a temporary directory to save uploaded files
    import tempfile
    import shutil

    temp_dir = tempfile.mkdtemp()

    try:
        for file in files:
            # 支持 .docx, .pdf, .txt, .md 文件
            supported_extensions = ('.docx', '.pdf', '.txt', '.md')
            if file.filename == '' or not file.filename.lower().endswith(supported_extensions) or file.filename.startswith('~$'):
                continue

            # Secure filename and save
            filename = file.filename
            # Handle paths in filename (for folder uploads)
            if '/' in filename:
                filename = filename.split('/')[-1]

            file_path = os.path.join(temp_dir, filename)
            file.save(file_path)

            char_count, status = get_word_count_unified(file_path)
            file_type = os.path.splitext(filename)[1].lower()  # 获取文件扩展名
            results.append({
                'filename': filename,
                'file_type': file_type,
                'char_count': char_count,
                'status': status
            })

    except Exception as e:
        return jsonify({'error': f'处理失败: {str(e)}'}), 500
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir)

    if not results:
        return jsonify({'error': '未找到有效的文件 (.docx, .pdf, .txt, .md)'}), 404

    # Sort by filename
    results.sort(key=lambda x: x['filename'])

    return jsonify({'results': results, 'count': len(results)})

@app.route('/api/export/excel', methods=['POST'])
def export_excel():
    data = request.json
    results = data.get('results', [])

    if not results:
        return jsonify({'error': '没有数据可导出'}), 400

    # Create Excel in memory
    output = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "字数统计报告"

    # Define headers
    headers = ["文件名", "文件类型", "字符数(不含空格)", "状态"]
    ws.append(headers)

    # Style headers
    header_font = Font(bold=True, color="FFFFFF", size=12)
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Add data
    for row in results:
        ws.append([row['filename'], row['file_type'], row['char_count'], row['status']])

    # Style data rows and adjust column width
    for row in ws.iter_rows(min_row=2, max_row=len(results) + 1):
        for cell in row:
            cell.alignment = Alignment(vertical="center")
            cell.border = thin_border
            if isinstance(cell.value, int):
                 cell.alignment = Alignment(horizontal="right", vertical="center")

    # Auto adjust column width
    for column_cells in ws.columns:
        length = max(len(str(cell.value)) for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = length * 1.2 + 2

    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name="字数统计报告.xlsx"
    )

@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    data = request.json
    results = data.get('results', [])

    if not results:
        return jsonify({'error': '没有数据可导出'}), 400

    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []

    # Register Chinese Font
    # Try to find a common Chinese font on macOS
    font_path = "/System/Library/Fonts/PingFang.ttc"
    font_name = "PingFang"
    try:
        if not os.path.exists(font_path):
             # Fallback to another common font if PingFang is not found
             font_path = "/System/Library/Fonts/STHeiti Light.ttc"
             font_name = "STHeiti"

        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont(font_name, font_path))
        else:
            # If no system font found, we might have an issue displaying Chinese.
            # For now, let's assume standard font (which won't show Chinese correctly) or try a relative path if user provided one.
            # But since we are on user's mac, these paths should likely exist.
            pass
    except Exception as e:
        print(f"Font registration failed: {e}")
        font_name = "Helvetica" # Fallback

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=18,
        alignment=1, # Center
        spaceAfter=20
    )
    normal_style = ParagraphStyle(
        'NormalStyle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10
    )

    # Title
    elements.append(Paragraph("文档字数统计报告", title_style))

    # Table Data
    table_data = [["文件名", "文件类型", "字符数", "状态"]]
    total_chars = 0
    for row in results:
        table_data.append([
            Paragraph(row['filename'], normal_style), # Wrap long filenames
            row['file_type'],
            str(row['char_count']),
            row['status']
        ])
        if isinstance(row['char_count'], int):
            total_chars += row['char_count']

    # Table Style
    table = Table(table_data, colWidths=[250, 60, 80, 80])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#4F81BD")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 20))

    # Summary
    elements.append(Paragraph(f"总文件数: {len(results)}", normal_style))
    elements.append(Paragraph(f"总字符数: {total_chars}", normal_style))

    doc.build(elements)
    output.seek(0)

    return send_file(
        output,
        mimetype="application/pdf",
        as_attachment=True,
        download_name="字数统计报告.pdf"
    )

if __name__ == '__main__':
    # Open browser automatically
    import webbrowser
    from threading import Timer
    def open_browser():
        webbrowser.open_new("http://127.0.0.1:5001")

    Timer(1, open_browser).start()
    app.run(debug=True, port=5001, use_reloader=False)
