import os
import csv
import docx
from flask import Flask, render_template, request, jsonify, send_file, Response
import io

app = Flask(__name__)

def get_word_count(file_path):
    """
    读取 docx 文件并统计字数 (统计逻辑：纯文本字符数，去除空格和换行)
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text_content = "".join(full_text)
        clean_text = text_content.replace(" ", "").replace("\n", "").replace("\t", "").replace("\r", "")
        return len(clean_text), "成功"
    except Exception as e:
        return 0, f"失败: {str(e)}"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/analyze', methods=['POST'])
def analyze():
    data = request.json
    folder_path = data.get('folder_path', '').strip()

    # Remove quotes if user dragged and dropped folder
    if (folder_path.startswith('"') and folder_path.endswith('"')) or (folder_path.startswith("'") and folder_path.endswith("'")):
        folder_path = folder_path[1:-1]

    if not folder_path or not os.path.isdir(folder_path):
        return jsonify({'error': '无效的文件夹路径，请检查后重试'}), 400

    docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.docx') and not f.startswith('~$')]

    if not docx_files:
        return jsonify({'error': '该文件夹下没有找到 .docx 文件'}), 404

    results = []
    for filename in docx_files:
        file_path = os.path.join(folder_path, filename)
        char_count, status = get_word_count(file_path)
        results.append({
            'filename': filename,
            'char_count': char_count,
            'status': status
        })

    # Sort by filename
    results.sort(key=lambda x: x['filename'])

    return jsonify({'results': results, 'count': len(results)})

@app.route('/api/analyze_upload', methods=['POST'])
def analyze_upload():
    if 'files[]' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    files = request.files.getlist('files[]')
    if not files:
        return jsonify({'error': '未选择文件'}), 400

    results = []

    # Create a temporary directory to save uploaded files
    import tempfile
    import shutil

    temp_dir = tempfile.mkdtemp()

    try:
        for file in files:
            if file.filename == '' or not file.filename.lower().endswith('.docx') or file.filename.startswith('~$'):
                continue

            # Secure filename and save
            filename = file.filename
            # Handle paths in filename (for folder uploads)
            if '/' in filename:
                filename = filename.split('/')[-1]

            file_path = os.path.join(temp_dir, filename)
            file.save(file_path)

            char_count, status = get_word_count(file_path)
            results.append({
                'filename': filename,
                'char_count': char_count,
                'status': status
            })

    except Exception as e:
        return jsonify({'error': f'处理失败: {str(e)}'}), 500
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir)

    if not results:
        return jsonify({'error': '未找到有效的 .docx 文件'}), 404

    # Sort by filename
    results.sort(key=lambda x: x['filename'])

    return jsonify({'results': results, 'count': len(results)})

@app.route('/api/export', methods=['POST'])
def export_csv():
    data = request.json
    results = data.get('results', [])

    if not results:
        return jsonify({'error': '没有数据可导出'}), 400

    # Create CSV in memory
    output = io.StringIO()
    # Add BOM for Excel compatibility
    output.write('\ufeff')
    writer = csv.writer(output)
    writer.writerow(["文件名", "字符数(不含空格)", "状态"])

    for row in results:
        writer.writerow([row['filename'], row['char_count'], row['status']])

    output.seek(0)

    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-disposition": "attachment; filename=字数统计结果.csv"}
    )

if __name__ == '__main__':
    # Open browser automatically
    import webbrowser
    from threading import Timer
    def open_browser():
        webbrowser.open_new("http://127.0.0.1:5001")

    Timer(1, open_browser).start()
    app.run(debug=True, port=5001, use_reloader=False)
