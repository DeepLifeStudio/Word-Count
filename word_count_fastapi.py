import os
import sys
import logging
from logging.handlers import RotatingFileHandler

# ============================================================================
# 日志初始化 - 必须在所有其他导入之前，以便捕获导入错误
# ============================================================================
log_file = os.path.join(os.path.expanduser("~"), "word_count_debug.log")
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        RotatingFileHandler(log_file, maxBytes=1024*1024, backupCount=3),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)
logger.info("=== Application starting ===")
logger.info(f"Python version: {sys.version}")
logger.info(f"Log file: {log_file}")

# ============================================================================
# 导入所有依赖 - 用 try-except 捕获导入错误
# ============================================================================
try:
    import io
    import tempfile
    import shutil
    import docx
    from typing import List
    from fastapi import FastAPI, Request, File, UploadFile, HTTPException, Form
    from fastapi.responses import StreamingResponse, HTMLResponse
    from fastapi.staticfiles import StaticFiles
    from fastapi.templating import Jinja2Templates
    from pydantic import BaseModel
    import pdfplumber
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    import markdown
    from bs4 import BeautifulSoup
    import chardet
    import re
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    logger.info("All imports successful")
except Exception as e:
    logger.critical(f"Import error: {e}", exc_info=True)
    sys.exit(1)

# ============================================================================
# PyInstaller 资源路径处理
# ============================================================================
def get_resource_path(relative_path):
    """
    获取资源文件的绝对路径，兼容开发环境和 PyInstaller 打包后的环境
    PyInstaller 会将资源文件解压到 sys._MEIPASS 临时目录
    """
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller 打包后的临时路径
        return os.path.join(sys._MEIPASS, relative_path)
    # 开发环境的相对路径
    return os.path.join(os.path.abspath("."), relative_path)

# ============================================================================
# FastAPI 应用实例
# ============================================================================
app = FastAPI(
    title="Word Count Pro API",
    description="文档字数统计工具 - FastAPI 高性能版本",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# 设置最大请求体大小为 100MB
app.max_file_size = 100 * 1024 * 1024

# 挂载静态文件和模板 - 使用资源路径处理函数
static_path = get_resource_path("static")
templates_path = get_resource_path("templates")

app.mount("/static", StaticFiles(directory=static_path), name="static")
templates = Jinja2Templates(directory=templates_path)

# ============================================================================
# Pydantic 数据模型
# ============================================================================
class AnalyzeRequest(BaseModel):
    folder_path: str

class FileResult(BaseModel):
    filename: str
    file_type: str
    char_count: int
    status: str

class ExportRequest(BaseModel):
    results: List[dict]

# ============================================================================
# 核心业务逻辑函数
# ============================================================================
def calculate_mixed_word_count(text):
    """
    计算混合字数:
    1. 中文字符(CJK): 每个字符计为1
    2. 英文/数字/其他: 以空格/标点分隔的单词计为1
    """
    if not text:
        return 0

    # 1. 统计中文字符 (CJK Unified Ideographs)
    cjk_pattern = re.compile(r'[\u4e00-\u9fff]')
    cjk_chars = cjk_pattern.findall(text)
    cjk_count = len(cjk_chars)

    # 2. 将中文字符替换为空格,然后统计剩下的英文单词
    no_cjk_text = cjk_pattern.sub(' ', text)
    english_words = no_cjk_text.split()
    english_count = len(english_words)

    return cjk_count + english_count

def get_word_count(file_path):
    """
    读取 docx 文件并统计字数 (统计逻辑:纯文本字符数,去除空格和换行)
    包含页眉和页脚内容
    """
    try:
        doc = docx.Document(file_path)
        full_text = []

        # 1. 正文内容
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 2. 页眉和页脚内容
        for section in doc.sections:
            # 页眉
            if section.header:
                for para in section.header.paragraphs:
                    full_text.append(para.text)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
            # 页脚
            if section.footer:
                for para in section.footer.paragraphs:
                    full_text.append(para.text)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)

        text_content = "\n".join(full_text)
        return calculate_mixed_word_count(text_content), "成功"
    except Exception as e:
        return 0, f"失败: {str(e)}"

def get_pdf_word_count(file_path):
    """
    读取 PDF 文件并统计字数 (统计逻辑:纯文本字符数,去除空格和换行)
    使用 pdfplumber 提高准确性
    """
    try:
        full_text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)

        text_content = "\n".join(full_text)
        return calculate_mixed_word_count(text_content), "成功"
    except Exception as e:
        return 0, f"失败: {str(e)}"

def get_txt_word_count(file_path):
    """
    读取 TXT 文件并统计字数 (统计逻辑:纯文本字符数,去除空格和换行)
    使用 chardet 自动检测编码
    """
    try:
        # 1. 读取二进制内容
        with open(file_path, 'rb') as f:
            raw_data = f.read()

        # 2. 检测编码
        result = chardet.detect(raw_data)
        encoding = result['encoding']

        # 3. 如果检测失败或置信度低,尝试常用编码
        if not encoding or result['confidence'] < 0.5:
            encodings_to_try = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        else:
            encodings_to_try = [encoding, 'utf-8', 'gbk'] # 优先尝试检测到的编码

        text_content = None
        for enc in encodings_to_try:
            try:
                text_content = raw_data.decode(enc)
                break
            except (UnicodeDecodeError, LookupError):
                continue

        if text_content is None:
             return 0, "失败: 无法识别文件编码"

        if text_content is None:
             return 0, "失败: 无法识别文件编码"

        return calculate_mixed_word_count(text_content), "成功"
    except Exception as e:
        return 0, f"失败: {str(e)}"

def get_md_word_count(file_path):
    """
    读取 Markdown 文件并统计字数
    逻辑: 将 Markdown 转换为 HTML, 然后提取纯文本, 去除 Markdown 语法符号
    """
    try:
        # Markdown 通常是 UTF-8, 但为了保险也检测一下
        with open(file_path, 'rb') as f:
            raw_data = f.read()

        result = chardet.detect(raw_data)
        encoding = result['encoding'] or 'utf-8'

        try:
            md_content = raw_data.decode(encoding)
        except:
            md_content = raw_data.decode('utf-8', errors='ignore')

        # 转换为 HTML
        html = markdown.markdown(md_content)

        # 提取纯文本
        soup = BeautifulSoup(html, 'html.parser')
        text_content = soup.get_text()

        # 提取纯文本
        soup = BeautifulSoup(html, 'html.parser')
        text_content = soup.get_text()

        return calculate_mixed_word_count(text_content), "成功"
    except Exception as e:
        return 0, f"失败: {str(e)}"

def get_word_count_unified(file_path):
    """
    统一的字数统计入口,根据文件扩展名分发到对应的处理函数
    支持 .docx, .pdf, .txt, .md 文件
    """
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == '.docx':
        return get_word_count(file_path)
    elif file_extension == '.pdf':
        return get_pdf_word_count(file_path)
    elif file_extension == '.txt':
        return get_txt_word_count(file_path)
    elif file_extension == '.md':
        return get_md_word_count(file_path)
    else:
        return 0, f"失败: 不支持的文件格式 {file_extension}"

# ============================================================================
# API 路由
# ============================================================================
@app.get('/', response_class=HTMLResponse)
async def index(request: Request):
    """主页 - 返回 Vue.js 单页应用"""
    return templates.TemplateResponse('index.html', {'request': request})

@app.post('/api/analyze')
async def analyze(data: AnalyzeRequest):
    """文件夹分析接口"""
    folder_path = data.folder_path.strip()

    # Remove quotes if user dragged and dropped folder
    if (folder_path.startswith('"') and folder_path.endswith('"')) or \
       (folder_path.startswith("'") and folder_path.endswith("'")):
        folder_path = folder_path[1:-1]

    if not folder_path or not os.path.isdir(folder_path):
        raise HTTPException(status_code=400, detail='无效的文件夹路径,请检查后重试')

    # 支持 .docx, .pdf, .txt, .md 文件
    supported_extensions = ('.docx', '.pdf', '.txt', '.md')
    supported_files = [f for f in os.listdir(folder_path)
                       if f.lower().endswith(supported_extensions) and not f.startswith('~$')]

    if not supported_files:
        raise HTTPException(status_code=404, detail='该文件夹下没有找到支持的文件 (.docx, .pdf, .txt, .md)')

    results = []
    for filename in supported_files:
        file_path = os.path.join(folder_path, filename)
        char_count, status = get_word_count_unified(file_path)
        file_type = os.path.splitext(filename)[1].lower()  # 获取文件扩展名
        results.append({
            'filename': filename,
            'file_type': file_type,
            'char_count': char_count,
            'status': status
        })

    # Sort by filename
    results.sort(key=lambda x: x['filename'])

    return {'results': results, 'count': len(results)}

@app.post('/api/analyze_upload')
async def analyze_upload(files: List[UploadFile] = File(..., alias='files[]')):
    """文件上传分析接口 - 接收前端发送的 files[] 字段"""
    if not files:
        raise HTTPException(status_code=400, detail='未上传文件')

    results = []

    # Create a temporary directory to save uploaded files
    temp_dir = tempfile.mkdtemp()

    try:
        for file in files:
            # 支持 .docx, .pdf, .txt, .md 文件
            supported_extensions = ('.docx', '.pdf', '.txt', '.md')
            if file.filename == '' or not file.filename.lower().endswith(supported_extensions) or file.filename.startswith('~$'):
                continue

            # Secure filename and save
            filename = file.filename
            # Handle paths in filename (for folder uploads)
            if '/' in filename:
                filename = filename.split('/')[-1]

            file_path = os.path.join(temp_dir, filename)

            # 异步保存文件
            content = await file.read()
            with open(file_path, 'wb') as f:
                f.write(content)

            char_count, status = get_word_count_unified(file_path)
            file_type = os.path.splitext(filename)[1].lower()  # 获取文件扩展名
            results.append({
                'filename': filename,
                'file_type': file_type,
                'char_count': char_count,
                'status': status
            })

    except Exception as e:
        shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail=f'处理失败: {str(e)}')
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir)

    if not results:
        raise HTTPException(status_code=404, detail='未找到有效的文件 (.docx, .pdf, .txt, .md)')

    # Sort by filename
    results.sort(key=lambda x: x['filename'])

    return {'results': results, 'count': len(results)}

@app.post('/api/export/excel')
async def export_excel(data: ExportRequest):
    """Excel 导出接口"""
    results = data.results

    if not results:
        raise HTTPException(status_code=400, detail='没有数据可导出')

    # Create Excel in memory
    output = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "字数统计报告"

    # Define headers
    headers = ["文件名", "文件类型", "字数(中字+英词)", "状态"]
    ws.append(headers)

    # Style headers
    header_font = Font(bold=True, color="FFFFFF", size=12)
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Add data
    for row in results:
        ws.append([row['filename'], row['file_type'], row['char_count'], row['status']])

    # Style data rows and adjust column width
    for row in ws.iter_rows(min_row=2, max_row=len(results) + 1):
        for cell in row:
            cell.alignment = Alignment(vertical="center")
            cell.border = thin_border
            if isinstance(cell.value, int):
                 cell.alignment = Alignment(horizontal="right", vertical="center")

    # Auto adjust column width
    for column_cells in ws.columns:
        length = max(len(str(cell.value)) for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = length * 1.2 + 2

    wb.save(output)
    output.seek(0)

    from urllib.parse import quote

    filename_encoded = quote('字数统计报告.xlsx')
    headers = {
        'Content-Disposition': f'attachment; filename="report.xlsx"; filename*=UTF-8\'\'\'{filename_encoded}'
    }

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers
    )

@app.post('/api/export/pdf')
async def export_pdf(data: ExportRequest):
    """PDF 导出接口"""
    results = data.results

    if not results:
        raise HTTPException(status_code=400, detail='没有数据可导出')

    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []

    # Register Chinese Font
    # Try to find a common Chinese font on macOS
    font_path = "/System/Library/Fonts/PingFang.ttc"
    font_name = "PingFang"
    try:
        if not os.path.exists(font_path):
             # Fallback to another common font if PingFang is not found
             font_path = "/System/Library/Fonts/STHeiti Light.ttc"
             font_name = "STHeiti"

        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont(font_name, font_path))
        else:
            # If no system font found, we might have an issue displaying Chinese.
            # For now, let's assume standard font (which won't show Chinese correctly) or try a relative path if user provided one.
            # But since we are on user's mac, these paths should likely exist.
            pass
    except Exception as e:
        print(f"Font registration failed: {e}")
        font_name = "Helvetica" # Fallback

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=18,
        alignment=1, # Center
        spaceAfter=20
    )
    normal_style = ParagraphStyle(
        'NormalStyle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10
    )

    # Title
    elements.append(Paragraph("文档字数统计报告", title_style))

    # Table Data
    table_data = [["文件名", "文件类型", "字符数", "状态"]]
    total_chars = 0
    for row in results:
        table_data.append([
            Paragraph(row['filename'], normal_style), # Wrap long filenames
            row['file_type'],
            str(row['char_count']),
            row['status']
        ])
        if isinstance(row['char_count'], int):
            total_chars += row['char_count']

    # Table Style
    table = Table(table_data, colWidths=[250, 60, 80, 80])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#4F81BD")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 20))

    # Summary
    elements.append(Paragraph(f"总文件数: {len(results)}", normal_style))
    elements.append(Paragraph(f"总字符数: {total_chars}", normal_style))

    doc.build(elements)
    output.seek(0)

    from urllib.parse import quote

    filename_encoded = quote('字数统计报告.pdf')
    headers = {
        'Content-Disposition': f'attachment; filename="report.pdf"; filename*=UTF-8\'\'\'{filename_encoded}'
    }

    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers=headers
    )

# ============================================================================
# 应用启动入口
# ============================================================================
if __name__ == '__main__':
    try:
        logger.info("Starting main application...")

        # Open browser automatically
        import webbrowser
        from threading import Timer
        import uvicorn

        def open_browser():
            url = "http://127.0.0.1:8000"
            logger.info(f"Attempting to open browser at {url}")
            try:
                webbrowser.open_new(url)
                logger.info("Browser open command sent")
            except Exception as e:
                logger.error(f"Failed to open browser: {e}")

        # 延迟 1.5 秒后打开浏览器,确保服务器已启动
        Timer(1.5, open_browser).start()

        # 启动 Uvicorn ASGI 服务器
        logger.info("Starting Uvicorn server...")
        uvicorn.run(
            app,
            host="127.0.0.1",
            port=8000,
            log_level="info",
            access_log=True
        )
    except Exception as e:
        logger.critical(f"Fatal error: {e}", exc_info=True)
        sys.exit(1)
